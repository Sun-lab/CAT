#!/usr/bin/env python
"""Write the combined single-positive table into a Word (.docx) document."""
import argparse
import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CATS = {
    "both (final cancer_reactive=True)": "Both (final cancer-reactive)",
    "single-positive: overall-high only (per_cell & ~cluster)": "Single-positive: overall-high only",
    "single-positive: in-cluster only (cluster & ~per_cell)": "Single-positive: cluster only",
    "neither": "Neither",
}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--indir", default=os.path.expanduser(
        "~/research/CAT/_revision/single_positive_combined"))
    ap.add_argument("--out", default=os.path.expanduser(
        "~/research/CAT/_revision/single_positive_combined/"
        "single_positive_table.docx"))
    args = ap.parse_args()

    summ = pd.read_csv(os.path.join(args.indir,
                                    "combined_single_positive_summary.csv"))
    frac = pd.read_csv(os.path.join(args.indir,
                                    "combined_cluster_fractions.csv"))

    # n (%) per study x cell_type x category
    summ["cell"] = summ.apply(
        lambda r: f"{int(r['n_cells']):,} ({r['pct_cells']:.1f}%)", axis=1)
    wide = summ.pivot_table(index=["study", "cell_type", "n_total_cells"],
                            columns="category", values="cell",
                            aggfunc="first").reset_index()
    ordered_cats = list(CATS.keys())
    wide = wide[["study", "cell_type", "n_total_cells"] + ordered_cats]

    # reactive-cluster counts (authoritative flag)
    rc = (frac.groupby(["study", "cell_type"])
          .agg(n_clusters=("cluster_id", "size"),
               n_reactive=("cluster_is_reactive", "sum"))
          .reset_index())
    rc_map = {(r.study, r.cell_type): (int(r.n_reactive), int(r.n_clusters))
              for r in rc.itertuples()}

    doc = Document()
    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading("Cancer-reactive T cells: overall vs cluster-based "
                    "classification", level=1)
    doc.add_paragraph(
        "A cell is called cancer-reactive only if it passes both criteria: "
        "(1) the overall per-cell signature call (pos_score > median and "
        "neg_score < median) and (2) the cluster call (the cell lies in a "
        "signature-gene Leiden cluster whose fraction of per-cell-positive "
        "cells exceeds 0.4). “Single-positive” cells pass only one "
        "criterion and are excluded by the final AND. Counts are over the "
        "TCR-resolved cells in each study.")

    headers = ["Study", "Cell type", "Total cells",
               "Both\n(final cancer-reactive)",
               "Single-positive:\noverall-high only",
               "Single-positive:\ncluster only", "Neither",
               "Reactive clusters\n(of total)"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for _, r in wide.iterrows():
        nr, nc = rc_map.get((r["study"], r["cell_type"]), ("", ""))
        row = table.add_row().cells
        vals = [r["study"], r["cell_type"], f"{int(r['n_total_cells']):,}"]
        vals += [r[c] for c in ordered_cats]
        vals += [f"{nr} / {nc}"]
        for j, v in enumerate(vals):
            row[j].text = str(v)
            for p in row[j].paragraphs:
                if j >= 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run(
        "Note: cluster reactive/non-reactive status follows the authoritative "
        "cancer_reactive_by_cluster flag assigned on the full data. Chow_2023 "
        "CD4 has no cluster exceeding the 0.4 threshold, so its final "
        "cancer-reactive fraction is 0.")
    run.italic = True
    run.font.size = Pt(8)

    doc.save(args.out)
    print(f"[written] {args.out}")


if __name__ == "__main__":
    main()
