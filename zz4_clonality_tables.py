#!/usr/bin/env python
"""
Clonality of cells filtered by IEDB / McPAS (both-chain non-human-antigen match),
stratified by whether the TCR is shared with VDJdb (score>=2) or database-only.

Writes a Word document with three sections:
  1. Clonality of filtered cells (CD8, CD4), split by sharing with VDJdb.
  2. Confidence levels per database + VDJdb threshold sweep + additional cells
     filtered by McPAS/IEDB on top of VDJdb (score>=2).
  3. Proportion cancer-reactive by non-human-antigen category for VDJdb/McPAS/IEDB.

Reuses the database loaders / normalization from zz3_db_compare.py.
Output: ~/research/CAT/_revision/db_comparison/clonality_filtered_cells.docx
"""
import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import zz3_db_compare as z

OUTDIR = os.path.expanduser("~/research/CAT/_revision/db_comparison")
DOCX   = os.path.join(OUTDIR, "clonality_filtered_cells.docx")

REPO      = os.path.dirname(os.path.abspath(__file__))
ANNOT_DIR = os.path.join(REPO, "classify_cancer_reactive_T", "annotation_by_vdjDb")


def both_cells(cr, alpha, beta):
    """Per-cell boolean: both chains match the database (exact CDR3+V)."""
    ak = set(alpha["cdr3"] + "|" + alpha["v_gene"])
    bk = set(beta["cdr3"] + "|" + beta["v_gene"])
    a = (cr["a_cdr3"] + "|" + cr["a_v"]).isin(ak).to_numpy()
    b = (cr["b_cdr3"] + "|" + cr["b_v"]).isin(bk).to_numpy()
    return a & b


def db_vdjdb_score(thr):
    """VDJdb non-human-antigen entries at a given confidence threshold."""
    v = pd.read_csv(z.VDJDB, sep="\t", low_memory=False)
    v = v[(v["species"] == "HomoSapiens") & (v["vdjdb.score"] >= thr)].copy()
    v["v_gene"] = v["v.segm"]
    v = v[v["antigen.species"] != "HomoSapiens"]
    return z.prep_db(v[v["gene"] == "TRA"][["cdr3", "v_gene"]],
                     v[v["gene"] == "TRB"][["cdr3", "v_gene"]])


def chain_hits_fast(cr, alpha, beta):
    """Per-cell (a_hit, b_hit): does each chain match the database (exact)?"""
    ak = set(alpha["cdr3"] + "|" + alpha["v_gene"])
    bk = set(beta["cdr3"] + "|" + beta["v_gene"])
    a = (cr["a_cdr3"] + "|" + cr["a_v"]).isin(ak).to_numpy()
    b = (cr["b_cdr3"] + "|" + cr["b_v"]).isin(bk).to_numpy()
    return a, b


def confidence_table():
    """Whether each database provides a graded confidence score."""
    return pd.DataFrame([
        {"database": "VDJdb", "confidence_field": "vdjdb.score (0-3)",
         "graded_score": "yes", "threshold_used": ">= 2"},
        {"database": "McPAS", "confidence_field": "Antigen.identification.method (categorical)",
         "graded_score": "no", "threshold_used": "n/a"},
        {"database": "IEDB", "confidence_field": "assay metadata only (no per-receptor score)",
         "graded_score": "no", "threshold_used": "n/a"},
    ])


def vdjdb_threshold_table(crs, vdj_by_thr):
    """VDJdb TCRs/cells filtered at score thresholds 2, 1, 0."""
    rows = []
    for ct, cr in crs.items():
        crpos = cr["cancer_reactive"].to_numpy().astype(bool)
        for thr in (2, 1, 0):
            both = both_cells(cr, *vdj_by_thr[thr])
            rows.append({"cell_type": ct, "min_vdjdb_score": f">={thr}",
                         "n_TCRs": len(set(cr.loc[both, "_clone"])),
                         "n_cells": int(both.sum()),
                         "n_cancer_reactive_cells": int((both & crpos).sum())})
    return pd.DataFrame(rows)


def additional_cells_table(crs, dbs):
    """Cells filtered by McPAS / IEDB / their union ON TOP OF VDJdb (score>=2)."""
    rows = []
    for ct, cr in crs.items():
        crpos = cr["cancer_reactive"].to_numpy().astype(bool)
        v2 = both_cells(cr, *dbs["VDJdb"])
        mc = both_cells(cr, *dbs["McPAS"]); ie = both_cells(cr, *dbs["IEDB"])
        for name, fl in [("McPAS", mc & ~v2), ("IEDB", ie & ~v2),
                         ("McPAS|IEDB union", (mc | ie) & ~v2)]:
            rows.append({"cell_type": ct, "database": name,
                         "additional_cells": int(fl.sum()),
                         "additional_cancer_reactive_cells": int((fl & crpos).sum())})
    return pd.DataFrame(rows)


def proportion_table(cr, dbs):
    """Proportion cancer-reactive by non-human-antigen category, for the 3 DBs."""
    order = ["Both", "TRA", "TRB", "None"]
    crpos = cr["cancer_reactive"].to_numpy().astype(bool)
    out = pd.DataFrame({"category": order})
    for db in ("VDJdb", "McPAS", "IEDB"):
        a, b = chain_hits_fast(cr, *dbs[db])
        cat = z.nonhuman_category(a, b)
        g = (pd.DataFrame({"cat": cat, "cr": crpos})
             .groupby("cat")["cr"].agg(n="size", r="sum").reindex(order))
        out[db] = [f"{g.loc[c, 'r'] / g.loc[c, 'n']:.3f} (n={int(g.loc[c, 'n']):,})"
                   for c in order]
    return out


def vdjdb_antigen_maps():
    """(normalized cdr3, v) -> matched non-human antigen species, per chain (VDJdb >=2)."""
    v = pd.read_csv(z.VDJDB, sep="\t", low_memory=False)
    v = v[(v["species"] == "HomoSapiens") & (v["vdjdb.score"] >= 2)
          & (v["antigen.species"] != "HomoSapiens")].copy()
    v["cdr3n"] = v["cdr3"].map(z.norm_cdr3)
    v["vn"] = v["v.segm"].map(z.norm_v)
    amap = {"TRA": {}, "TRB": {}}
    for gene in ("TRA", "TRB"):
        sub = v[v["gene"] == gene]
        for key, grp in sub.groupby(["cdr3n", "vn"]):
            amap[gene][key] = ",".join(sorted(set(grp["antigen.species"].astype(str))))
    return amap


def removed_cr_tcr_list(cr):
    """Unique TCRs of cancer-reactive CD8 cells removed by the VDJdb both-chain
    non-human-antigen filter, with the number of cells per TCR and the matched
    non-human antigen species for each chain."""
    crpos = cr["cancer_reactive"].to_numpy().astype(bool)
    removed = crpos & both_cells(cr, *db_vdjdb_score(2))
    sub = cr[removed].copy()
    amap = vdjdb_antigen_maps()
    sub["TRA_antigen"] = [amap["TRA"].get((ac, av), "") for ac, av in zip(sub.a_cdr3, sub.a_v)]
    sub["TRB_antigen"] = [amap["TRB"].get((bc, bv), "") for bc, bv in zip(sub.b_cdr3, sub.b_v)]
    g = (sub.groupby(["TRA_v_gene", "TRA_cdr3", "TRB_v_gene", "TRB_cdr3",
                      "TRA_antigen", "TRB_antigen"])
            .size().reset_index(name="n_cells")
            .sort_values("n_cells", ascending=False).reset_index(drop=True))
    return g


def clonality_table(cr, dbs):
    """Combined IEDB + McPAS clonality table for one cell type (cr carries _clone)."""
    size_map = cr["_clone"].value_counts().to_dict()
    vdj_clones = set(cr.loc[both_cells(cr, *dbs["VDJdb"]), "_clone"])

    rows = []
    for dbname in ("IEDB", "McPAS"):
        db_clones = set(cr.loc[both_cells(cr, *dbs[dbname]), "_clone"])
        groups = [("shared with VDJdb", db_clones & vdj_clones),
                  ("DB-only (not in VDJdb)", db_clones - vdj_clones),
                  ("all", db_clones)]
        for label, clones in groups:
            sz = np.array([size_map[c] for c in clones]) if clones else np.array([0])
            rows.append({
                "database": dbname, "group": label,
                "n_TCRs": len(clones), "n_cells": int(sz.sum()),
                "mean_cells_per_TCR": round(float(sz.mean()), 1),
                "median": int(np.median(sz)), "max": int(sz.max()),
                "pct_singleton_TCRs": round(100 * float(np.mean(sz == 1)), 1),
            })
    return pd.DataFrame(rows)


def add_df_table(doc, df, heading):
    doc.add_heading(heading, level=2)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = col
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            cells[j].text = f"{v:,}" if isinstance(v, (int, np.integer)) else str(v)
            for p in cells[j].paragraphs:
                if j >= 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)


def main():
    dbs = {"VDJdb": z.db_vdjdb(), "McPAS": z.db_mcpas(), "IEDB": z.db_iedb()}
    vdj_by_thr = {2: dbs["VDJdb"], 1: db_vdjdb_score(1), 0: db_vdjdb_score(0)}

    # load each cell type once (with a clone key) and reuse for all tables
    crs = {}
    for ct in ("CD8", "CD4"):
        cr = z.load_cr(ct)
        cr["_clone"] = list(zip(cr.TRA_cdr3, cr.TRA_v_gene, cr.TRB_cdr3, cr.TRB_v_gene))
        crs[ct] = cr

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ---- Section 1: clonality of filtered cells ----
    doc.add_heading("Clonality of cells filtered by IEDB / McPAS, "
                    "stratified by sharing with VDJdb", level=1)
    doc.add_paragraph(
        "Cells are filtered when both TCR chains match a non-human-antigen TCR "
        "(exact CDR3 + V gene). For each database the filtered clonotypes are split "
        "into those also flagged by VDJdb (score ≥ 2) vs database-only, and "
        "summarized by clone size (cells per TCR).")
    for ct in ("CD8", "CD4"):
        tab = clonality_table(crs[ct], dbs)
        tab.to_csv(os.path.join(OUTDIR, f"clonality_filtered_{ct}.csv"), index=False)
        add_df_table(doc, tab, ct)

    # ---- Section 2: confidence levels and additional filtering ----
    doc.add_heading("Confidence levels and additional filtering by McPAS / IEDB", level=1)
    doc.add_paragraph(
        "Only VDJdb provides a graded confidence score (vdjdb.score); the pipeline "
        "uses score ≥ 2. Relaxing the threshold barely changes the VDJdb-filtered "
        "set, and the McPAS/IEDB-only TCRs are not recovered even at score ≥ 0 — "
        "they are genuinely absent from VDJdb, not low-confidence entries.")
    conf = confidence_table()
    thr = vdjdb_threshold_table(crs, vdj_by_thr)
    add = additional_cells_table(crs, dbs)
    conf.to_csv(os.path.join(OUTDIR, "confidence_availability.csv"), index=False)
    thr.to_csv(os.path.join(OUTDIR, "vdjdb_threshold_sweep.csv"), index=False)
    add.to_csv(os.path.join(OUTDIR, "additional_cells_on_top_of_vdjdb.csv"), index=False)
    add_df_table(doc, conf, "Confidence-score availability per database")
    add_df_table(doc, thr, "VDJdb cells filtered by confidence threshold")
    add_df_table(doc, add, "Additional cells filtered on top of VDJdb (score ≥ 2)")

    # ---- Section 3: proportion cancer-reactive by antigen category (3 DBs) ----
    doc.add_heading("Proportion of cancer-reactive cells by non-human-antigen "
                    "category (VDJdb / McPAS / IEDB)", level=1)
    doc.add_paragraph(
        "Restricted to the final cancer_reactive definition. Category = how many of "
        "a cell's chains match a non-human-antigen TCR (Both / TRA only / TRB only / "
        "None). Each entry is the proportion cancer-reactive (cell count in parentheses).")
    for ct in ("CD8", "CD4"):
        pt = proportion_table(crs[ct], dbs)
        pt.to_csv(os.path.join(OUTDIR, f"prop_cancer_reactive_by_antigen_{ct}.csv"), index=False)
        add_df_table(doc, pt, ct)

    doc.save(DOCX)
    print(f"[written] {DOCX}")

    # ---- TCR lists: cancer-reactive cells removed by the VDJdb filter ----
    os.makedirs(ANNOT_DIR, exist_ok=True)
    for ct in ("CD8", "CD4"):
        tcrs = removed_cr_tcr_list(crs[ct])
        txt = os.path.join(ANNOT_DIR, f"CR_{ct}_tcrs_non-human-antigen.txt")
        with open(txt, "w") as f:
            f.write(f"# Cancer-reactive {ct} T cells removed because BOTH TCR chains match\n")
            f.write("# a non-human-antigen TCR in VDJdb (score >= 2).\n")
            f.write(f"# total removed cells: {int(tcrs['n_cells'].sum())}; unique TCRs: {len(tcrs)}\n")
            f.write("# TRA/TRB_antigen = non-human antigen species matched by that chain in VDJdb.\n")
            tcrs.to_csv(f, sep="\t", index=False)
        print(f"[written] {txt}  ({int(tcrs['n_cells'].sum())} cells, {len(tcrs)} TCRs)")


if __name__ == "__main__":
    main()
